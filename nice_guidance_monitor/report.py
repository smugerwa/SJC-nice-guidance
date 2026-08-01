from __future__ import annotations

import re
from pathlib import Path

# House style: the clinic's brand palette (skinjointclinic.co.uk) - navy and
# gold, matching the Cliniko audit reports.
INK = "101820"
NAVY = "071727"
MUTED = "596675"
TABLE_HEADER_FILL = "E9DFD0"
RULE_GREY = "e4e7eb"


def _letterhead_name(report: dict, config: dict) -> str:
    return config.get("letterhead_name") or report.get("practice_name", "")


def build_markdown_report(report: dict) -> str:
    source_label = report.get("source_label", "NICE")
    report_title = report.get("report_title", f"{source_label} Guidance Monthly Review")
    period_label = report.get("period_label", "Reporting month")
    prepared_by = report.get("prepared_by", f"{source_label} Guidance Monitoring Agent")
    included = [i for i in report["items_reviewed"] if i.get("included")]
    excluded = [i for i in report["items_reviewed"] if not i.get("included")]
    included = sorted(included, key=lambda item: item.get("relevance", {}).get("score", 0), reverse=True)
    high_min = report.get("thresholds", {}).get("high_relevance_min_score", 4)
    high = [i for i in included if i.get("relevance", {}).get("score", 0) >= high_min]

    lines = [
        f"# {report_title} - {report['month_label']} - {report['practice_name']}",
        "",
        f"**{period_label}:** {report['month_label']}",
        f"**Date generated:** {report['date_generated']}",
        f"**Prepared by:** {prepared_by}",
        f"**Reviewed by:** {report.get('reviewer') or '[INSERT NAME/ROLE]'}",
        "",
        "## Executive summary",
        "",
        f"- {source_label} items reviewed: {len(report['items_reviewed'])}",
        f"- Included in detailed review: {len(included)}",
        f"- Excluded or appendix only: {len(excluded)}",
        f"- High or very high relevance: {len(high)}",
        "",
    ]
    relevance_label = report.get("relevance_label", "Primary care relevance")

    clinically_relevant = [i for i in included if i.get("relevance", {}).get("score", 0) >= 3]
    lines += ["### Key points for clinical meeting", ""]
    for item in clinically_relevant[:6]:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        lines.append(f"- **{_item_reference(ident)} - {ident.get('title')}:** {brief.get('what_changed', '').strip()}")
    if not clinically_relevant:
        lines.append("- No clinically relevant updates identified for routine primary care.")
    lines.append("")

    lines += [
        "## Action dashboard",
        "",
        f"| {source_label} item | Relevance | What to do | Meeting question |",
        "| --- | ---: | --- | --- |",
    ]
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        lines.append(
            f"| [{_item_reference(ident)} - {ident.get('title')}]({ident.get('url')}) "
            f"| {item.get('relevance', {}).get('score', '')} "
            f"| {brief.get('suggested_action', '')} "
            f"| {brief.get('meeting_discussion', '')} |"
        )
    lines.append("")

    lines += ["## Clinical Update Briefs", ""]
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        lines += [
            f"### {_item_reference(ident)} - {ident.get('title')}",
            "",
            f"- **Type:** {ident.get('guidance_type', '')}",
            f"- **Date:** {ident.get('publication_or_update_date', '')}",
            f"- **Status:** {ident.get('status', '')}",
            f"- **{relevance_label}:** {item.get('relevance', {}).get('score', '')}/5 - {item.get('relevance', {}).get('rationale', '')}",
            f"- **{source_label} source:** {ident.get('url', '')}",
            "",
            "#### What changed or matters",
            "",
            brief.get("what_changed", ""),
            "",
            "#### Key takeaways for clinicians",
            "",
        ]
        takeaways = brief.get("key_takeaways") or item.get("key_clinical_points", [])[:5]
        lines += [f"- {point}" for point in takeaways]
        lines += [
            "",
            "#### Practice implication",
            "",
            brief.get("practice_implication", ""),
            "",
            "#### Suggested meeting discussion",
            "",
            f"- {brief.get('meeting_discussion', '')}",
            "",
            "#### Suggested action",
            "",
            f"- {brief.get('suggested_action', '')}",
            "",
        ]

    lines += ["## Items for clinical meeting", ""]
    discussion = []
    decisions = []
    awareness = []
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        score = item.get("relevance", {}).get("score", 0)
        label = f"{_item_reference(ident)} - {ident.get('title')}: {item.get('clinical_brief', {}).get('meeting_discussion', '')}"
        if score >= high_min:
            decisions.append(label)
        elif score >= 3:
            discussion.append(label)
        else:
            awareness.append(label)
    lines += ["**Items requiring discussion:**"] + ([f"- {x}" for x in discussion] or ["- None"])
    lines += ["", "**Items requiring decision:**"] + ([f"- {x}" for x in decisions] or ["- None"])
    lines += ["", "**Items for awareness only:**"] + ([f"- {x}" for x in awareness] or ["- None"])
    lines.append("")

    low_relevance = [i for i in included if i.get("relevance", {}).get("score", 0) < 3]
    lines += [f"## Appendix A: Low-Relevance Or Excluded {source_label} Items", ""]
    lines += [f"| Title | {source_label} reference | URL | Reason excluded |", "| --- | --- | --- | --- |"]
    for item in excluded:
        ident = item["guidance_identification"]
        lines.append(f"| {ident.get('title')} | {_item_reference(ident)} | {ident.get('url')} | {item.get('exclusion_reason', '')} |")
    for item in low_relevance:
        ident = item["guidance_identification"]
        lines.append(f"| {ident.get('title')} | {_item_reference(ident)} | {ident.get('url')} | Low relevance; awareness only. |")
    if not excluded and not low_relevance:
        lines.append("| None |  |  |  |")
    lines.append("")

    lines += [f"## Appendix B: Main {source_label} Sources", ""]
    for item in report["items_reviewed"]:
        ident = item.get("guidance_identification", {})
        if ident.get("url"):
            lines.append(f"- {_item_reference(ident)} - {ident.get('title')}: {ident.get('url')}")
    if report.get("failures"):
        lines += ["", "## Source retrieval failures", ""]
        lines += [f"- {failure}" for failure in report["failures"]]
    lines += [
        "",
        "## Actions completed",
        "",
        "_To be completed by the practice team: record any actions taken in response to the items above, with dates - evidence of implementation of changes in practice (CQC)._",
        "",
        "| Date | Action taken | Evidence / notes |",
        "| --- | --- | --- |",
        "|  |  |  |",
        "|  |  |  |",
        "|  |  |  |",
    ]
    lines += ["", "_Clinical governance document - for internal use._", ""]
    lines += ["", "_Full linked-source extraction is retained in the JSON source log for audit, but omitted from this meeting brief for readability._", ""]
    return "\n".join(lines)


def build_docx_report(report: dict, path: Path, config: dict) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        return
    source_label = report.get("source_label", "NICE")
    report_title = report.get("report_title", f"{source_label} Guidance Monthly Review")
    period_label = report.get("period_label", "Reporting month")
    prepared_by = report.get("prepared_by", f"{source_label} Guidance Monitoring Agent")

    template = config.get("headed_paper_template_docx")
    doc = Document(template) if template and Path(template).exists() else Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    _setup_docx_styles(doc)

    ink = RGBColor.from_string(INK)
    navy = RGBColor.from_string(NAVY)
    muted = RGBColor.from_string(MUTED)
    letterhead = _letterhead_name(report, config)
    for doc_section in doc.sections:
        _letterhead_header(doc_section, letterhead, f"{report_title} — {report['month_label']}", navy, muted)
        _page_number_footer(doc_section, f"{letterhead} — {report_title}, {report['month_label']}", muted)

    included = [i for i in report["items_reviewed"] if i.get("included")]
    excluded = [i for i in report["items_reviewed"] if not i.get("included")]
    included = sorted(included, key=lambda item: item.get("relevance", {}).get("score", 0), reverse=True)
    high_min = report.get("thresholds", {}).get("high_relevance_min_score", 4)
    clinically_relevant = [i for i in included if i.get("relevance", {}).get("score", 0) >= 3]
    high = [i for i in clinically_relevant if i.get("relevance", {}).get("score", 0) >= high_min]

    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report_title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Calibri"
    run.font.color.rgb = navy
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{report['month_label']} - {report['practice_name']}")
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.color.rgb = muted
    doc.add_paragraph()

    meta = _add_table(doc, ["Field", "Detail"], [
        [period_label, report["month_label"]],
        ["Date generated", report["date_generated"]],
        ["Prepared by", prepared_by],
        ["Reviewed by", report.get("reviewer") or "[INSERT NAME/ROLE]"],
    ], [2.0, 4.6])
    _shade_header(meta)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    relevance_label = report.get("relevance_label", "Primary care relevance")
    doc.add_heading("Executive Summary", 1)
    _add_bullet(doc, f"{source_label} items reviewed: {len(report['items_reviewed'])}")
    _add_bullet(doc, f"Included in clinical brief: {len(clinically_relevant)}")
    _add_bullet(doc, f"Excluded or appendix only: {len(excluded)}")
    _add_bullet(doc, f"High or very high relevance: {len(high)}")

    doc.add_heading("Key Points For Clinical Meeting", 2)
    for item in clinically_relevant[:6]:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        p = _add_bullet(doc, "")
        _add_inline_markup(p, f"**{_item_reference(ident)} - {ident.get('title')}:** {brief.get('what_changed', '').strip()}")

    doc.add_heading("Action Dashboard", 1)
    action_rows = []
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        action_rows.append([
            f"{_item_reference(ident)} - {ident.get('title')}",
            str(item.get("relevance", {}).get("score", "")),
            brief.get("suggested_action", ""),
            brief.get("meeting_discussion", ""),
        ])
    action_table = _add_table(doc, [f"{source_label} item", "Score", "What to do", "Meeting question"], action_rows, [2.15, 0.55, 2.25, 2.4])
    _shade_header(action_table)

    doc.add_heading("Clinical Update Briefs", 1)
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        brief = item.get("clinical_brief", {})
        doc.add_heading(f"{_item_reference(ident)} - {ident.get('title')}", 2)
        _add_label_value(doc, "Type", ident.get("guidance_type", ""))
        _add_label_value(doc, "Date", ident.get("publication_or_update_date", ""))
        _add_label_value(doc, "Status", ident.get("status", ""))
        _add_label_value(doc, relevance_label, f"{item.get('relevance', {}).get('score', '')}/5 - {item.get('relevance', {}).get('rationale', '')}")
        _add_label_value(doc, f"{source_label} source", ident.get("url", ""))

        doc.add_heading("What Changed Or Matters", 3)
        doc.add_paragraph(_clean_docx_text(brief.get("what_changed", "")))
        doc.add_heading("Key Takeaways For Clinicians", 3)
        for point in brief.get("key_takeaways") or item.get("key_clinical_points", [])[:5]:
            _add_bullet(doc, _clean_docx_text(point))
        doc.add_heading("Practice Implication", 3)
        doc.add_paragraph(_clean_docx_text(brief.get("practice_implication", "")))
        doc.add_heading("Suggested Meeting Discussion", 3)
        _add_bullet(doc, _clean_docx_text(brief.get("meeting_discussion", "")))
        doc.add_heading("Suggested Action", 3)
        _add_bullet(doc, _clean_docx_text(brief.get("suggested_action", "")))

    doc.add_heading("Items For Clinical Meeting", 1)
    meeting_rows = []
    for item in clinically_relevant:
        ident = item["guidance_identification"]
        score = item.get("relevance", {}).get("score", 0)
        category = "Decision" if score >= high_min else "Discussion"
        meeting_rows.append([category, f"{_item_reference(ident)} - {ident.get('title')}", item.get("clinical_brief", {}).get("meeting_discussion", "")])
    meeting_table = _add_table(doc, ["Type", "Item", "Meeting prompt"], meeting_rows, [1.0, 2.6, 3.8])
    _shade_header(meeting_table)

    low_relevance = [i for i in included if i.get("relevance", {}).get("score", 0) < 3]
    doc.add_heading(f"Appendix A: Low-Relevance Or Excluded {source_label} Items", 1)
    appendix_rows = []
    for item in excluded:
        ident = item["guidance_identification"]
        appendix_rows.append([_item_reference(ident), ident.get("title", ""), item.get("exclusion_reason", "")])
    for item in low_relevance:
        ident = item["guidance_identification"]
        appendix_rows.append([_item_reference(ident), ident.get("title", ""), "Low relevance; awareness only."])
    appendix_table = _add_table(doc, ["Reference", "Title", "Reason"], appendix_rows or [["None", "", ""]], [0.85, 3.25, 3.3])
    _shade_header(appendix_table)

    doc.add_heading(f"Appendix B: Main {source_label} Sources", 1)
    source_rows = []
    for item in report["items_reviewed"]:
        ident = item.get("guidance_identification", {})
        if ident.get("url"):
            source_rows.append([_item_reference(ident), ident.get("title", ""), ident.get("url", "")])
    source_table = _add_table(doc, ["Reference", "Guidance", "URL"], source_rows, [0.85, 3.25, 3.3])
    _shade_header(source_table)

    # Completed by hand after each review - never written by the automation.
    doc.add_heading("Actions Completed", 1)
    guidance_note = doc.add_paragraph()
    guidance_run = guidance_note.add_run(
        "To be completed by the practice team: record any actions taken in response to the "
        "items above, with dates - evidence of implementation of changes in practice (CQC)."
    )
    guidance_run.italic = True
    guidance_run.font.size = Pt(8.5)
    guidance_run.font.name = "Calibri"
    guidance_run.font.color.rgb = muted
    actions_table = _add_table(
        doc,
        ["Date", "Action taken", "Evidence / notes"],
        [["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""]],
        [1.0, 3.0, 2.65],
    )
    _shade_header(actions_table)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Full linked-source extraction is retained in the JSON source log for audit, but omitted from this meeting brief for readability.")
    r.italic = True
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = muted

    _finalize_docx_typography(doc)
    doc.save(path)


def _item_reference(ident: dict) -> str:
    return ident.get("source_reference") or ident.get("nice_reference") or ident.get("reference") or "MHRA"


def _setup_docx_styles(doc) -> None:
    from docx.shared import Pt, RGBColor

    ink = RGBColor.from_string(INK)
    navy = RGBColor.from_string(NAVY)
    for style_name, size, color, bold in [
        ("Normal", 12, ink, False),
        ("Heading 1", 16, navy, True),
        ("Heading 2", 13, navy, True),
        ("Heading 3", 12, navy, True),
    ]:
        style = _get_style_case_insensitive(doc, style_name)
        if style:
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = bold
    _set_uk_language(doc)


def _set_uk_language(doc) -> None:
    """Mark the document as UK English so Word proofs against en-GB."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    style = _get_style_case_insensitive(doc, "Normal")
    if not style:
        return
    rpr = style.element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-GB")


def _letterhead_header(section, letterhead: str, right_text: str, navy, muted) -> None:
    """Clinic letterhead on every page: name in brand navy on the left, the
    report title and period muted on the right, with a rule underneath."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.65), WD_TAB_ALIGNMENT.RIGHT)
    _paragraph_rule(p, edge="bottom")
    name_run = p.add_run(letterhead)
    name_run.bold = True
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(9.5)
    name_run.font.color.rgb = navy
    right_run = p.add_run(f"\t{right_text}")
    right_run.font.name = "Calibri"
    right_run.font.size = Pt(7.5)
    right_run.font.color.rgb = muted


def _page_number_footer(section, left_text: str, muted) -> None:
    """'{clinic} — {report}, {period}' left, 'Page N/M' right, with a rule
    above - the footer every page of the report carries."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.65), WD_TAB_ALIGNMENT.RIGHT)
    _paragraph_rule(p, edge="top")

    def styled(text=""):
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(7.5)
        run.font.color.rgb = muted
        return run

    styled(f"{left_text}\tPage ")
    # page numbers are Word fields, so they render per page
    _add_field(styled(), "PAGE")
    styled("/")
    _add_field(styled(), "NUMPAGES")


def _paragraph_rule(paragraph, edge: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    border = OxmlElement("w:pBdr")
    line = OxmlElement(f"w:{edge}")
    for attr, val in (("w:val", "single"), ("w:sz", "4"), ("w:space", "6"), ("w:color", RULE_GREY)):
        line.set(qn(attr), val)
    border.append(line)
    paragraph._p.get_or_add_pPr().append(border)


def _add_field(run, instruction: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for element in ("begin", None, "end"):
        if element is None:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = instruction
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), element)
            run._r.append(fld)


def _get_style_case_insensitive(doc, style_name: str):
    for style in doc.styles:
        if style.name.lower() == style_name.lower():
            return style
    return None


def _finalize_docx_typography(doc) -> None:
    from docx.shared import Pt, RGBColor

    navy = RGBColor.from_string(NAVY)
    body = RGBColor.from_string(INK)
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name.lower()
        for run in paragraph.runs:
            if not run.text:
                continue
            if not run.font.name:
                run.font.name = "Calibri"
            if style_name.startswith("heading"):
                if style_name == "heading 1":
                    run.font.size = Pt(16)
                elif style_name == "heading 2":
                    run.font.size = Pt(13)
                else:
                    run.font.size = Pt(12)
                run.font.color.rgb = navy
                run.bold = True
            elif not run.font.size:
                run.font.size = Pt(12)
                run.font.color.rgb = body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.font.name:
                            run.font.name = "Calibri"


def _add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = _clean_docx_text(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = _clean_docx_text(str(value or ""))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def _shade_header(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_HEADER_FILL)
        tc_pr.append(shading)


def _add_bullet(doc, text: str):
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
    except KeyError:
        paragraph = doc.add_paragraph()
        paragraph.add_run("• ")
    if text:
        _add_inline_markup(paragraph, text)
    return paragraph


def _add_label_value(doc, label: str, value: str):
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    _add_inline_markup(paragraph, value)
    return paragraph


def _add_inline_markup(paragraph, text: str) -> None:
    text = _clean_docx_text(text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _clean_docx_text(text: str) -> str:
    replacements = {
        "â€“": "-",
        "â€”": "-",
        "â€™": "'",
        "â€˜": "'",
        "â€œ": '"',
        "â€": '"',
        "Â": "",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.replace("**", "").strip()
